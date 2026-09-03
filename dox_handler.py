import os
import subprocess
import shutil
from docx import Document
from docx.shared import Inches

libreoffice = "C:\\Program Files\\LibreOffice\\program\\soffice.exe"

"""Each function opens document ,performs related operations and saves.
Document closes by itself after being saved, so no need to close it.
Also there is no such method as 'doc.close()' in docx library"""
def create_docx(save_path):
    """Create a new, blank .docx and save it to save_path."""
    doc = Document()
    doc.save(save_path)
    return save_path


def copy_docx(open_path, save_path):
    """Open an existing .docx from open_path and save a copy to save_path."""
    # doc = Document(open_path)
    # doc.save(save_path)
    # return save_path
    if not os.path.exists(open_path):
        raise FileNotFoundError(
            f"Template DOCX not found: {open_path}"
        )

    output_dir = os.path.dirname(
        os.path.abspath(save_path)     #absolute path gives complete path
    )

    os.makedirs(output_dir, exist_ok=True)

    shutil.copy2(open_path, save_path)

    return save_path
    

def write_text_to_docx(open_path, save_path, text):
    """Open a .docx, add a paragraph of text, save the result."""
    if not os.path.exists(open_path):
        raise FileNotFoundError(
            f"DOCX file not found: {open_path}"
        )

    doc = Document(open_path)
    doc.add_paragraph(str(text))
    doc.save(save_path)
    return save_path
    

def add_picture_to_docx(open_path, save_path, image_path, width_inches=4):
    """Open a .docx, add a picture, save the result."""
    doc = Document(open_path)
    doc.add_picture(image_path, width=Inches(float(width_inches)))
    doc.save(save_path)    #object document closes after saving.
    return save_path


def save_docx(document, save_path):
    """Saves an existing Document object to save_path."""
    document.save(save_path)   #save existing Document object ,writes the document to disk.
    return save_path

def delete_file(file_path):
    """
    Delete a file if it exists.
    """

    if os.path.exists(file_path):
        os.remove(file_path)

    return True

def convert_docx_to_pdf(open_path, save_path):
    """Convert the .docx at open_path into a .pdf saved at save_path."""
    open_path = os.path.abspath(open_path)
    save_path = os.path.abspath(save_path)
    outdir = os.path.dirname(save_path)
    os.makedirs(outdir, exist_ok=True)

    if not os.path.exists(libreoffice):
        raise FileNotFoundError(
            f"LibreOffice not found: {libreoffice}"
        )

    if not os.path.exists(open_path):
        raise FileNotFoundError(
            f"DOCX file not found: {open_path}"
        )
    try:
        subprocess.run(
            [libreoffice, "--headless", "--convert-to", "pdf", "--outdir", outdir, open_path],
            check=True,
        )

        # LibreOffice names the PDF after the input file -- rename to save_path if different
        produced_path = os.path.join(
            outdir, os.path.splitext(os.path.basename(open_path))[0] + ".pdf"
        )
        if produced_path != save_path:
            if os.path.exists(save_path):  #to avoid failure if a file already exists
                os.remove(save_path)
            os.replace(produced_path, save_path)
        return save_path
    finally:
        if  os.path.exists(save_path):          #if pdf file is created
            if os.path.exists(open_path):       #if docx file exists 
                os.remove(open_path)            #remove docx file

